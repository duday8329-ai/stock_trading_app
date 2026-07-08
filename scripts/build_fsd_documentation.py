from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
    ListFlowable,
    ListItem,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "fsd-documentation"
OUT.mkdir(parents=True, exist_ok=True)

PROJECT_TITLE = "STOCK TRADING APP"
REPO = "https://github.com/duday8329-ai/stock_trading_app"
DEMO = "https://sb-stocks-docker.onrender.com"

TEAM = [
    ("UDAY DONIKELA", "teamLead"),
    ("Kunduru Leela Durga Prasad", "member"),
    ("Edula Pranathi Chandana Reddy", "member"),
    ("EPURI VENKATA LAKSHMI VARSHITHA", "member"),
    ("Kuruva Eswaraiah", "member"),
]

API_ROWS = [
    ("POST", "/api/auth/register", "Register a new user account."),
    ("POST", "/api/auth/login", "Authenticate user and return JWT token."),
    ("GET", "/api/auth/me", "Return logged-in user profile."),
    ("GET", "/api/stocks/search?q=", "Search stocks by ticker or company name."),
    ("GET", "/api/stocks/quote/:ticker", "Return current quote for a ticker."),
    ("GET", "/api/stocks/history/:ticker?range=", "Return historical price data."),
    ("POST", "/api/trade/buy", "Buy shares using virtual cash."),
    ("POST", "/api/trade/sell", "Sell shares from existing holdings."),
    ("GET", "/api/portfolio", "Return cash, holdings, values, and gain/loss."),
    ("GET", "/api/transactions", "Return trade transaction history."),
    ("GET/POST", "/api/watchlist", "View or add watchlist tickers."),
    ("DELETE", "/api/watchlist/:ticker", "Remove ticker from watchlist."),
    ("GET/POST", "/api/admin/stocks", "Admin stock listing and creation."),
    ("PUT/DELETE", "/api/admin/stocks/:id", "Admin stock update and delete."),
]


def bullet_docx(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def number_docx(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_docx_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    for row_values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_values):
            row[idx].text = value
    return table


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Full Stack Development with MERN")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Functional Specification Document")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string("2E74B5")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(PROJECT_TITLE).bold = True
    doc.add_paragraph()

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(f"Project Title: {PROJECT_TITLE}")
    doc.add_paragraph("Team Members:")
    add_docx_table(doc, ["Name", "Role"], TEAM)

    doc.add_heading("2. Project Overview", level=1)
    doc.add_paragraph(
        "Purpose: SB Stocks is a full-stack stock trading simulation platform that lets users practice buying and selling stocks with virtual funds. The goal is to provide a safe learning environment for stock market practice without real brokerage, payment, or real-money integration."
    )
    doc.add_paragraph("Features:")
    bullet_docx(
        doc,
        [
            "User registration, login, logout, and JWT-protected routes.",
            "Dashboard with cash balance, holdings, total portfolio value, and gain/loss.",
            "Stock search, quote display, and historical chart.",
            "Virtual buy/sell trading engine with validation.",
            "Transaction history and watchlist management.",
            "Admin stock management panel.",
            "Responsive frontend deployed with backend on Render.",
        ],
    )

    doc.add_heading("3. Architecture", level=1)
    doc.add_paragraph("Frontend: React with functional components, hooks, React Router, Tailwind CSS, and Recharts.")
    doc.add_paragraph("Backend: Node.js and Express.js with modular routes, controllers, services, middleware, and error handling.")
    doc.add_paragraph("Database: MongoDB Atlas with Mongoose schemas for User, Holding, Transaction, Watchlist, and Stock.")
    doc.add_paragraph("Deployment: Docker-based Render web service serving both the React build and Express API.")

    doc.add_heading("4. Setup Instructions", level=1)
    doc.add_paragraph("Prerequisites:")
    bullet_docx(doc, ["Node.js 20+", "npm", "MongoDB Atlas account", "Finnhub API key", "Git"])
    doc.add_paragraph("Installation:")
    number_docx(
        doc,
        [
            f"Clone the repository: {REPO}",
            "Copy server/.env.example to server/.env.",
            "Add MONGO_URI, JWT_SECRET, FINNHUB_API_KEY, STARTING_CASH, and MARKET_CACHE_TTL_SECONDS.",
            "Install dependencies using npm run install:all.",
            "Seed demo data using npm run seed.",
            "Start local development using npm run dev.",
        ],
    )

    doc.add_heading("5. Folder Structure", level=1)
    doc.add_paragraph("Client: React frontend source is organized into components, pages, hooks, and services.")
    bullet_docx(doc, ["client/src/components", "client/src/pages", "client/src/hooks", "client/src/services"])
    doc.add_paragraph("Server: Express backend source is organized into config, controllers, middleware, models, routes, and services.")
    bullet_docx(doc, ["server/src/config", "server/src/controllers", "server/src/middleware", "server/src/models", "server/src/routes", "server/src/services"])

    doc.add_heading("6. Running the Application", level=1)
    bullet_docx(
        doc,
        [
            "Install all dependencies: npm run install:all",
            "Seed demo data: npm run seed",
            "Run both client and server: npm run dev",
            "Frontend local URL: http://localhost:5173",
            "Backend local URL: http://localhost:5000",
            f"Production live demo: {DEMO}",
        ],
    )

    doc.add_heading("7. API Documentation", level=1)
    add_docx_table(doc, ["Method", "Endpoint", "Description"], API_ROWS)
    doc.add_paragraph("Example protected response: /api/portfolio returns cash balance, holdings, holding value, total portfolio value, and overall gain/loss.")

    doc.add_heading("8. Authentication", level=1)
    doc.add_paragraph(
        "Authentication is handled using JWT tokens. User passwords are hashed with bcrypt before being stored in MongoDB. After login, the backend returns a signed JWT. Protected API requests include the token in the Authorization header as a Bearer token."
    )
    bullet_docx(
        doc,
        [
            "requireAuth middleware validates user sessions.",
            "requireAdmin middleware protects admin-only stock management APIs.",
            "Auth routes are rate-limited to reduce brute-force attempts.",
            "Secrets are stored in environment variables and are not committed to Git.",
        ],
    )

    doc.add_heading("9. User Interface", level=1)
    bullet_docx(
        doc,
        [
            "Login and Register pages for user access.",
            "Dashboard page for cash, holdings, portfolio value, and gain/loss.",
            "Stock Detail page with quote, chart, and buy/sell form.",
            "Transactions page for chronological trade history.",
            "Watchlist page for saved tickers.",
            "Profile page for account information.",
            "Admin Stocks page for admin CRUD operations.",
        ],
    )

    doc.add_heading("10. Testing", level=1)
    bullet_docx(
        doc,
        [
            "Authentication flow tested through register, login, and protected /me endpoint.",
            "Trading validation tested for insufficient cash and insufficient shares.",
            "Portfolio updates tested after buy and sell transactions.",
            "Transaction history tested after simulated orders.",
            "Deployment health tested using /api/health on Render.",
        ],
    )

    doc.add_heading("11. Screenshots or Demo", level=1)
    doc.add_paragraph(f"Live demo: {DEMO}")
    doc.add_paragraph(f"GitHub repository: {REPO}")
    doc.add_paragraph("Suggested screenshots for submission: Login, Register, Dashboard, Stock Detail, Buy/Sell form, Transactions, Watchlist, Admin Stocks, Render deployment page, and MongoDB Atlas collections.")

    doc.add_heading("12. Known Issues", level=1)
    bullet_docx(
        doc,
        [
            "Render free service can spin down during inactivity and may take time to wake up.",
            "Free-tier market data can be delayed or rate-limited.",
            "Fallback mock market data is used when the Finnhub API key is unavailable.",
        ],
    )

    doc.add_heading("13. Future Enhancements", level=1)
    bullet_docx(
        doc,
        [
            "Leaderboards across users.",
            "Multiple virtual portfolios per user.",
            "Price alerts.",
            "Dark mode.",
            "More detailed analytics and portfolio charts.",
        ],
    )

    path = OUT / "FSD_STOCK_TRADING_APP.docx"
    doc.save(path)
    return path


def pdf_list(items, styles):
    return ListFlowable([ListItem(Paragraph(item, styles["Body"])) for item in items], bulletType="bullet", leftIndent=18)


def build_pdf():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleBlue", parent=styles["Title"], textColor=colors.HexColor("#0B2545"), fontSize=20, leading=24, alignment=1))
    styles.add(ParagraphStyle(name="SubtitleBlue", parent=styles["Heading2"], textColor=colors.HexColor("#2E74B5"), alignment=1))
    styles.add(ParagraphStyle(name="H1Blue", parent=styles["Heading1"], textColor=colors.HexColor("#2E74B5"), fontSize=15, leading=18, spaceBefore=12, spaceAfter=6))
    styles.add(ParagraphStyle(name="Body", parent=styles["BodyText"], fontSize=10.5, leading=14, spaceAfter=6))

    story = [
        Paragraph("Full Stack Development with MERN", styles["TitleBlue"]),
        Paragraph("Functional Specification Document", styles["SubtitleBlue"]),
        Paragraph(PROJECT_TITLE, styles["SubtitleBlue"]),
        Spacer(1, 0.15 * inch),
        Paragraph("1. Introduction", styles["H1Blue"]),
        Paragraph(f"Project Title: {PROJECT_TITLE}", styles["Body"]),
    ]
    team_table = Table([["Name", "Role"], *TEAM], colWidths=[4.3 * inch, 1.7 * inch])
    team_table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")), ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story += [team_table, Spacer(1, 0.1 * inch)]

    sections = [
        ("2. Project Overview", ["Purpose: SB Stocks is a full-stack stock trading simulation platform for practicing stock trades with virtual funds.", "Features include authentication, dashboard, stock search, quote/history, buy/sell simulation, transactions, watchlist, admin stock management, and responsive UI."]),
        ("3. Architecture", ["Frontend: React, React Router, Tailwind CSS, Recharts.", "Backend: Node.js, Express.js, controllers, routes, middleware, and services.", "Database: MongoDB Atlas with Mongoose models for User, Holding, Transaction, Watchlist, and Stock."]),
        ("4. Setup Instructions", ["Prerequisites: Node.js 20+, npm, MongoDB Atlas, Finnhub API key, Git.", "Installation: clone repo, copy server/.env.example, set variables, run npm run install:all, npm run seed, and npm run dev."]),
        ("5. Folder Structure", ["Client: client/src/components, pages, hooks, services.", "Server: server/src/config, controllers, middleware, models, routes, services."]),
        ("6. Running the Application", [f"Local frontend: http://localhost:5173", "Local backend: http://localhost:5000", f"Live demo: {DEMO}"]),
    ]
    for heading, paragraphs in sections:
        story.append(Paragraph(heading, styles["H1Blue"]))
        for paragraph in paragraphs:
            story.append(Paragraph(paragraph, styles["Body"]))

    story.append(Paragraph("7. API Documentation", styles["H1Blue"]))
    api_table = Table([["Method", "Endpoint", "Description"], *API_ROWS], colWidths=[0.8 * inch, 2.1 * inch, 3.1 * inch], repeatRows=1)
    api_table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.35, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")), ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 8), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story.append(api_table)

    more_sections = [
        ("8. Authentication", ["JWT tokens are used for sessions. Passwords are hashed with bcrypt. Protected routes use requireAuth, and admin APIs use requireAdmin."]),
        ("9. User Interface", ["Pages include Login, Register, Dashboard, Stock Detail, Transactions, Profile, Watchlist, and Admin Stocks."]),
        ("10. Testing", ["Tested authentication, protected routes, trading validation, portfolio updates, transaction history, and Render health checks."]),
        ("11. Screenshots or Demo", [f"GitHub Repository: {REPO}", f"Live Demo: {DEMO}", "Screenshots can include login, dashboard, stock detail, transaction history, admin panel, Render, and MongoDB Atlas."]),
        ("12. Known Issues", ["Render free service can spin down with inactivity. Free-tier market data can be delayed or rate-limited."]),
        ("13. Future Enhancements", ["Leaderboards, multiple virtual portfolios, price alerts, dark mode, and advanced portfolio analytics."]),
    ]
    for heading, paragraphs in more_sections:
        story.append(Paragraph(heading, styles["H1Blue"]))
        for paragraph in paragraphs:
            story.append(Paragraph(paragraph, styles["Body"]))

    path = OUT / "FSD_STOCK_TRADING_APP.pdf"
    pdf = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=0.75 * inch, rightMargin=0.75 * inch, topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    pdf.build(story)
    return path


def main():
    docx_path = build_docx()
    pdf_path = build_pdf()
    print(docx_path)
    print(pdf_path)


if __name__ == "__main__":
    main()
