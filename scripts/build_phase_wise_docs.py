from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "phase-wise-submission"
REPO = "https://github.com/duday8329-ai/stock_trading_app"
DEMO = "https://sb-stocks-docker.onrender.com"


def set_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("555555")

    doc.add_paragraph()


def add_meta(doc, phase, progress="100%"):
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Project", "SB Stocks - Paper Trading Platform"),
        ("Phase", phase),
        ("Progress", progress),
        ("GitHub Repository", REPO),
        ("Live Demo", DEMO),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
        row.cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_story_table(doc, stories):
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Story", "Duration", "Status", "Description", "Resources"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
        table.rows[0].cells[idx].paragraphs[0].runs[0].bold = True
    for story in stories:
        row = table.add_row().cells
        row[0].text = story["story"]
        row[1].text = story["duration"]
        row[2].text = "Completed"
        row[3].text = story["description"]
        row[4].text = story["resources"]


def save_doc(filename, title, phase, sections):
    doc = Document()
    set_styles(doc)
    add_title(doc, title, "Completed phase document for mentor review")
    add_meta(doc, phase)
    for section in sections:
        doc.add_heading(section["heading"], level=1)
        if "paragraphs" in section:
            for paragraph in section["paragraphs"]:
                doc.add_paragraph(paragraph)
        if "bullets" in section:
            add_bullets(doc, section["bullets"])
        if "numbers" in section:
            add_numbered(doc, section["numbers"])
        if "stories" in section:
            add_story_table(doc, section["stories"])
        if "table" in section:
            headers, rows = section["table"]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
                table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            for values in rows:
                cells = table.add_row().cells
                for i, value in enumerate(values):
                    cells[i].text = value
        doc.add_paragraph()
    path = OUT / filename
    doc.save(path)
    return path


def main():
    OUT.mkdir(parents=True, exist_ok=True)

    common_resources = f"Repository: {REPO}\nDemo: {DEMO}"

    save_doc(
        "01_Define_Problem_Statements.docx",
        "Define Problem Statements",
        "Brainstorming and Ideation",
        [
            {
                "heading": "Problem Statement",
                "paragraphs": [
                    "New stock market learners need a safe platform to practice buying and selling stocks without risking real money. Many beginners find live brokerage platforms intimidating because mistakes can create real financial loss.",
                    "SB Stocks solves this problem with a simulation-only paper trading platform that uses virtual funds, realistic stock data, portfolio tracking, and transaction history.",
                ],
            },
            {
                "heading": "Target Users",
                "bullets": [
                    "Students learning MERN stack development and stock market workflows.",
                    "Beginners who want to practice trading without real money.",
                    "Mentors/reviewers evaluating full-stack project implementation.",
                ],
            },
            {
                "heading": "Success Criteria",
                "bullets": [
                    "Users can register and log in securely.",
                    "Users can search stocks and view quote/history data.",
                    "Users can buy/sell stocks using virtual funds.",
                    "Portfolio and transaction history update correctly.",
                    "The application is deployed and accessible online.",
                ],
            },
        ],
    )

    save_doc(
        "02_Empathy_Map_Canvas.docx",
        "Empathy Map Canvas",
        "Brainstorming and Ideation",
        [
            {
                "heading": "User Says",
                "bullets": [
                    "I want to learn stock trading without losing money.",
                    "I need a simple dashboard to understand my holdings.",
                    "I want to see my trade history clearly.",
                ],
            },
            {
                "heading": "User Thinks",
                "bullets": [
                    "Real trading platforms are risky for beginners.",
                    "A simulation can help me understand buy/sell decisions.",
                    "Clear validation messages help me learn mistakes safely.",
                ],
            },
            {
                "heading": "User Does",
                "bullets": [
                    "Registers an account.",
                    "Searches for stocks.",
                    "Places simulated buy/sell trades.",
                    "Reviews portfolio and transactions.",
                ],
            },
            {
                "heading": "User Feels",
                "bullets": [
                    "More confident practicing in a risk-free environment.",
                    "In control because the app shows cash, holdings, and gains/losses.",
                    "Safe because no real brokerage or payment integration exists.",
                ],
            },
        ],
    )

    save_doc(
        "03_Brainstorming_Idea_Generation_Prioritization.docx",
        "Brainstorming, Idea Generation, and Prioritization",
        "Brainstorming and Ideation",
        [
            {
                "heading": "Ideas Considered",
                "bullets": [
                    "Paper trading with virtual cash.",
                    "Watchlist management.",
                    "Admin stock management.",
                    "Leaderboards and multi-portfolio support.",
                    "Price alerts and dark mode.",
                ],
            },
            {
                "heading": "Prioritized MVP",
                "numbers": [
                    "Authentication with JWT and bcrypt.",
                    "Dashboard with portfolio metrics.",
                    "Stock search and detail chart.",
                    "Buy/sell simulation engine.",
                    "Transaction history.",
                    "Render deployment.",
                ],
            },
            {
                "heading": "Deferred Items",
                "bullets": [
                    "Leaderboards.",
                    "Multiple portfolios per user.",
                    "Price alerts.",
                    "Dark mode.",
                ],
            },
        ],
    )

    save_doc(
        "04_Solution_Requirements.docx",
        "Solution Requirements",
        "Requirement Analysis",
        [
            {
                "heading": "Functional Requirements",
                "bullets": [
                    "User registration, login, logout, and JWT-protected routes.",
                    "Dashboard with cash balance, holdings, total portfolio value, and gain/loss.",
                    "Stock search, quote, and historical chart.",
                    "Buy/sell validation for cash and share quantity.",
                    "Transaction history with price, quantity, type, and timestamp.",
                    "Watchlist and admin stock management.",
                ],
            },
            {
                "heading": "Non-Functional Requirements",
                "bullets": [
                    "Environment variables for all secrets.",
                    "Helmet, CORS, auth rate limiting, and error handling.",
                    "Responsive frontend layout.",
                    "Docker deployment on Render.",
                    "MongoDB Atlas persistence.",
                ],
            },
        ],
    )

    save_doc(
        "05_Technology_Stack.docx",
        "Technology Stack",
        "Requirement Analysis",
        [
            {
                "heading": "Selected Stack",
                "table": (
                    ["Layer", "Technology", "Purpose"],
                    [
                        ["Frontend", "React, Vite, React Router, Tailwind CSS", "Responsive single-page application"],
                        ["Charts", "Recharts", "Historical price visualization"],
                        ["Backend", "Node.js, Express.js", "REST API and business logic"],
                        ["Database", "MongoDB Atlas, Mongoose", "Cloud persistence and schemas"],
                        ["Security", "JWT, bcrypt, Helmet, CORS", "Authentication and protection"],
                        ["Market Data", "Finnhub API", "Stock search, quote, and history"],
                        ["Deployment", "Docker, Render", "Production hosting"],
                    ],
                ),
            }
        ],
    )

    save_doc(
        "06_Data_Flow_Diagrams_and_User_Stories.docx",
        "Data Flow Diagrams and User Stories",
        "Requirement Analysis",
        [
            {
                "heading": "High-Level Data Flow",
                "numbers": [
                    "User opens the React frontend.",
                    "Frontend sends API requests to Express backend.",
                    "Backend validates JWT and request payloads.",
                    "Backend reads/writes MongoDB through Mongoose models.",
                    "Backend calls Finnhub or fallback market-data service.",
                    "Frontend updates dashboard, stock page, and transaction screens.",
                ],
            },
            {
                "heading": "User Stories",
                "stories": [
                    {
                        "story": "Register/Login",
                        "duration": "45m",
                        "description": "As a user, I can create an account and log in securely.",
                        "resources": common_resources,
                    },
                    {
                        "story": "Trade Stocks",
                        "duration": "2h",
                        "description": "As a user, I can buy and sell stocks with virtual funds.",
                        "resources": common_resources,
                    },
                    {
                        "story": "View Portfolio",
                        "duration": "1h",
                        "description": "As a user, I can view holdings, cash, portfolio value, and gain/loss.",
                        "resources": common_resources,
                    },
                ],
            },
        ],
    )

    save_doc(
        "07_Project_Planning.docx",
        "Project Planning",
        "Project Planning Phase",
        [
            {
                "heading": "Phase Plan",
                "table": (
                    ["Phase", "Duration", "Status", "Deliverable"],
                    [
                        ["Project Architecture", "1h 0m", "100%", "Architecture docs, ER diagram, user flow, MVC pattern"],
                        ["Setup and Configuration", "30m", "100%", "Client/server setup, env example, deployment config"],
                        ["Backend Development", "8h 0m", "100%", "Auth, trading APIs, portfolio, watchlist, admin APIs"],
                        ["Database Development", "2h 0m", "100%", "MongoDB connection and Mongoose models"],
                        ["Frontend Development", "8h 0m", "100%", "React pages, routes, responsive UI"],
                        ["Project Execution", "30m", "100%", "Local steps, Render deployment, live demo"],
                    ],
                ),
            },
            {
                "heading": "Risks and Mitigation",
                "bullets": [
                    "Market API rate limits are reduced through caching and fallback data.",
                    "Secrets are protected through environment variables and ignored `.env` files.",
                    "Deployment environment issues are handled with Docker.",
                ],
            },
        ],
    )

    save_doc(
        "08_Problem_Solution_Fit.docx",
        "Problem Solution Fit",
        "Project Design Phase",
        [
            {
                "heading": "Problem",
                "paragraphs": [
                    "Beginner traders need a realistic but safe platform to practice trading concepts without using real funds or brokerage accounts."
                ],
            },
            {
                "heading": "Solution Fit",
                "bullets": [
                    "Virtual cash removes real-money risk.",
                    "Market quote/history data makes practice realistic.",
                    "Portfolio and transaction tracking show learning outcomes.",
                    "Role-based admin panel demonstrates full-stack CRUD and security.",
                ],
            },
        ],
    )

    save_doc(
        "09_Proposed_Solution.docx",
        "Proposed Solution",
        "Project Design Phase",
        [
            {
                "heading": "Solution Overview",
                "paragraphs": [
                    "SB Stocks is a full-stack paper-trading platform where users receive virtual funds and can simulate stock trades using live or delayed market data. The system tracks cash balance, holdings, portfolio value, gain/loss, watchlists, and transaction history."
                ],
            },
            {
                "heading": "Modules",
                "bullets": [
                    "Authentication module.",
                    "Market data module.",
                    "Trading and portfolio module.",
                    "Transaction history module.",
                    "Watchlist module.",
                    "Admin stock management module.",
                ],
            },
        ],
    )

    save_doc(
        "10_Solution_Architecture.docx",
        "Solution Architecture",
        "Project Design Phase",
        [
            {
                "heading": "Architecture Layers",
                "bullets": [
                    "React frontend handles pages, forms, charts, and API calls.",
                    "Express backend handles REST APIs, validation, security, and business logic.",
                    "MongoDB Atlas stores users, holdings, stocks, watchlists, and transactions.",
                    "Finnhub provides stock search, quote, and historical candle data.",
                    "Render Docker service hosts the production frontend and API.",
                ],
            },
            {
                "heading": "MVC Mapping",
                "table": (
                    ["Pattern Area", "Project Location"],
                    [
                        ["Models", "server/src/models"],
                        ["Views", "client/src/pages and client/src/components"],
                        ["Controllers", "server/src/controllers"],
                        ["Routes", "server/src/routes"],
                        ["Middleware", "server/src/middleware"],
                        ["Services", "server/src/services"],
                    ],
                ),
            },
        ],
    )

    save_doc(
        "11_User_Acceptance_Testing_FSD.docx",
        "User Acceptance Testing and FSD",
        "Project Development",
        [
            {
                "heading": "Feature Specification",
                "bullets": [
                    "Users can register and log in with JWT authentication.",
                    "Users can search stocks and view quote/history information.",
                    "Users can buy and sell shares using virtual funds only.",
                    "Portfolio and transactions update after every trade.",
                    "Admins can manage stock records.",
                ],
            },
            {
                "heading": "Acceptance Tests",
                "table": (
                    ["Test Case", "Expected Result", "Status"],
                    [
                        ["Register new user", "Account is created with virtual cash", "Pass"],
                        ["Login user", "JWT session is created", "Pass"],
                        ["Buy with insufficient cash", "Order is rejected", "Pass"],
                        ["Sell more shares than owned", "Order is rejected", "Pass"],
                        ["Buy valid quantity", "Holding and transaction are created", "Pass"],
                        ["Open live app", "Render service loads successfully", "Pass"],
                    ],
                ),
            },
            {
                "heading": "Final Output",
                "bullets": [
                    f"GitHub Repository: {REPO}",
                    f"Live Demo: {DEMO}",
                    "Overall tracker progress: 100%",
                ],
            },
        ],
    )

    print(f"Created {len(list(OUT.glob('*.docx')))} DOCX files in {OUT}")


if __name__ == "__main__":
    main()
